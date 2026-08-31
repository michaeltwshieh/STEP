from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/michaelshieh/Desktop/Claude Projects/STEP Exam/KnowledgeBase/Company-Law-Exam-Bot-Routing-v2-Test")
OUT = ROOT / "output" / "docx" / "Routing-v2-Question-1-Step-by-Step-Guide-Cantonese.docx"

BLUE = "214E78"
DARK = "17324D"
INK = "23313F"
MUTED = "5D6B78"
PALE_BLUE = "E8F1F8"
PALE_GOLD = "FFF4D6"
PALE_GREEN = "EAF5EE"
LIGHT = "F4F6F8"
WHITE = "FFFFFF"
RED = "8C2F39"
GRID = "C7D1DB"

PRESET = {
    "name": "compact_reference_guide",
    "page": {"width": 8.5, "height": 11.0, "margins": 1.0, "header": 0.492, "footer": 0.492},
    "body": {"font": "Arial Unicode MS", "east_asia": "Arial Unicode MS", "size": 10.5, "after": 6, "line": 1.25},
    "h1": {"size": 16, "before": 18, "after": 10, "color": BLUE},
    "h2": {"size": 13, "before": 14, "after": 7, "color": BLUE},
    "h3": {"size": 12, "before": 10, "after": 5, "color": DARK},
    "table": {"width_dxa": 9360, "indent_dxa": 120, "cell_top": 90, "cell_bottom": 90, "cell_start": 120, "cell_end": 120},
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size=6):
    tblPr = table._tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_in):
    widths_dxa = [round(w * 1440) for w in widths_in]
    assert sum(widths_dxa) == PRESET["table"]["width_dxa"], (widths_in, widths_dxa, sum(widths_dxa))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(PRESET["table"]["width_dxa"]))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(PRESET["table"]["indent_dxa"]))
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        trPr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[i])
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[i]))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_run_font(run, size=None, bold=None, color=None, italic=None, mono=False):
    name = "Courier New" if mono else PRESET["body"]["font"]
    east = PRESET["body"]["east_asia"]
    run.font.name = name
    run._element.get_or_add_rPr()
    fonts = run._element.rPr.rFonts
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), east)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(p, after=6, before=0, line=1.25, keep=False):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_together = keep
    for run in p.runs:
        set_run_font(run, size=PRESET["body"]["size"], color=INK)


def add_text(doc, text="", after=6, before=0, bold_prefix=None, keep=False):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, size=PRESET["body"]["size"], color=DARK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=PRESET["body"]["size"], color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=PRESET["body"]["size"], color=INK)
    style_paragraph(p, after=after, before=before, keep=keep)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    token = PRESET[f"h{level}"]
    set_run_font(r, size=token["size"], bold=True, color=token["color"])
    p.paragraph_format.space_before = Pt(token["before"])
    p.paragraph_format.space_after = Pt(token["after"])
    p.paragraph_format.line_spacing = 1.0
    return p


def add_file_line(doc, file_name, why, next_files, result=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run("睇咗：")
    set_run_font(r, bold=True, color=BLUE, size=10.5)
    r = p.add_run(file_name)
    set_run_font(r, mono=True, color=DARK, size=9.4)
    r = p.add_run("\n點解：")
    set_run_font(r, bold=True, color=BLUE, size=10.5)
    r = p.add_run(why)
    set_run_font(r, color=INK, size=10.5)
    r = p.add_run("\n再指去：")
    set_run_font(r, bold=True, color=BLUE, size=10.5)
    r = p.add_run(next_files)
    set_run_font(r, mono=True, color=DARK, size=9.4)
    if result:
        r = p.add_run("\n今步產出：")
        set_run_font(r, bold=True, color=BLUE, size=10.5)
        r = p.add_run(result)
        set_run_font(r, color=INK, size=10.5)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, border=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.line_spacing = 1.18
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(label + "  ")
    set_run_font(r, bold=True, color=border, size=10.5)
    r = p.add_run(text)
    set_run_font(r, color=INK, size=10.5)
    return p


def add_table(doc, headers, rows, widths, header_fill=BLUE, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    mark_repeat_header(hdr)
    for i, text in enumerate(headers):
        set_cell_shading(hdr.cells[i], header_fill)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(text)
        set_run_font(r, size=font_size, bold=True, color=WHITE)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if ridx % 2 == 1:
                set_cell_shading(cells[i], LIGHT)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK, mono=(i == 1 and ("." in str(value) or "/" in str(value))))
    set_table_geometry(table, widths)
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" of ")
    set_run_font(run, size=8.5, color=MUTED)
    fld2 = OxmlElement("w:fldSimple")
    fld2.set(qn("w:instr"), "NUMPAGES")
    paragraph._p.append(fld2)


def setup_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(PRESET["page"]["width"])
    sec.page_height = Inches(PRESET["page"]["height"])
    sec.top_margin = Inches(PRESET["page"]["margins"])
    sec.bottom_margin = Inches(PRESET["page"]["margins"])
    sec.left_margin = Inches(PRESET["page"]["margins"])
    sec.right_margin = Inches(PRESET["page"]["margins"])
    sec.header_distance = Inches(PRESET["page"]["header"])
    sec.footer_distance = Inches(PRESET["page"]["footer"])
    sec.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = PRESET["body"]["font"]
    normal._element.rPr.rFonts.set(qn("w:ascii"), PRESET["body"]["font"])
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), PRESET["body"]["font"])
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), PRESET["body"]["east_asia"])
    normal.font.size = Pt(PRESET["body"]["size"])
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(PRESET["body"]["after"])
    normal.paragraph_format.line_spacing = PRESET["body"]["line"]

    for i in (1, 2, 3):
        st = doc.styles[f"Heading {i}"]
        token = PRESET[f"h{i}"]
        st.font.name = PRESET["body"]["font"]
        st._element.rPr.rFonts.set(qn("w:eastAsia"), PRESET["body"]["east_asia"])
        st.font.size = Pt(token["size"])
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(token["color"])
        st.paragraph_format.space_before = Pt(token["before"])
        st.paragraph_format.space_after = Pt(token["after"])
        st.paragraph_format.keep_with_next = True

    header = sec.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    r = hp.add_run("Routing v2 | Question 1 worked example")
    set_run_font(r, size=8.2, color=MUTED)
    r = hp.add_run("\tSTEP Company Law test package")
    set_run_font(r, size=8.2, color=MUTED)

    fp = sec.footer.paragraphs[0]
    add_page_number(fp)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(78)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("WORKED ROUTING GUIDE")
    set_run_font(r, size=10.5, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Routing v2 點樣由 Question 1\n推算到答案")
    set_run_font(r, size=28, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("逐步睇清楚：睇咩 file、點解睇、下一步去邊個 file")
    set_run_font(r, size=13, color=BLUE)

    add_callout(
        doc,
        "名稱確認",
        "測試包內正式名稱係 Routing v2，冇獨立一套叫 Routing 3-2。本文按你嘅用語，將『Routing 3-2』理解為呢套 Routing v2 流程。",
        fill=PALE_GOLD,
        border="8A6A18",
    )

    add_text(doc, "例子：STEP Advanced Certificate in Company Law and Practice，Specimen Examination Paper 2，Question 1。", after=4)
    add_text(doc, "目的：解釋答案點樣由本地 course files、考試附件同 deterministic RoutePlan 一步一步產生；本文唔係重新作答。", after=4)
    add_text(doc, "狀態：Routing v2 test package。呢份文件唔代表 live workflow 已經啟用，亦唔係 GO 決定。", after=4)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared as a transparent worked example")
    set_run_font(r, size=9.5, italic=True, color=MUTED)
    doc.add_page_break()


def build_document():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    add_cover(doc)

    add_heading(doc, "1. 先講最簡單版本：成條 routing 路徑", 1)
    add_callout(
        doc,
        "一句講晒",
        "先由考試題目抽 facts，再用 Content.md 將每個 issue 指去相應 Module／Appendix；routing-core.md 將呢啲 routes 鎖定、過濾、hash 同驗證；驗證合格後，section-b.md 先批准寫成 submit-ready 答案。",
        fill=PALE_GREEN,
        border="2B6A45",
    )

    overview_rows = [
        ("0", "入口規則", "AGENTS.md → CLAUDE.md", "知道只可以用本資料夾及用邊套流程"),
        ("1", "讀題", "Exam PDF：Q1 + exam Appendix 1", "抽出角色、要求、20 marks、11 組 questionnaire facts"),
        ("2", "分類", "CLAUDE.md + routing-core.md", "Q1 = PROSE；唔係 MCQ，亦唔係 DRAFTING"),
        ("3", "搵 source", "Content.md", "由每個 issue 加總出 Modules 2/3/5/6/7/9/10/11 + Appendices"),
        ("4", "鎖定路線", "routing-core.md", "六個 locks、facts ledger、second pass、relevance gate"),
        ("5", "封 source", "RoutePlan JSON", "freeze allowlist；只可開列明嘅 files，逐個記 SHA-256"),
        ("6", "驗證", "schema + validate_route_plan.py", "15 個 invariants 全 PASS，status = VALID"),
        ("7", "寫答案", "section-b.md", "輸出 SUBMIT THIS + DO NOT SUBMIT check panel"),
    ]
    add_table(doc, ["Stage", "做咩", "主要 file", "今步產出"], overview_rows, [0.55, 1.0, 2.55, 2.4], font_size=8.4)

    add_heading(doc, "2. 三種 file，角色完全唔同", 1)
    type_rows = [
        ("流程指令", "AGENTS.md、CLAUDE.md、routing-core.md、section-b.md、schema、validator", "話畀系統知『點做』；本身唔係法律權威。"),
        ("法律內容", "Course-Manual-Module-*.md、Appendix-*.md", "提供規則、程序、例子同 precedent；答案嘅法律內容由呢度嚟。"),
        ("考試輸入", "Specimen Paper 2 PDF，Q1 同 exam Appendix 1", "決定 facts、問題、marks 同 requested deliverable。"),
    ]
    add_table(doc, ["類別", "例子", "作用"], type_rows, [1.0, 2.9, 2.6], font_size=8.8)
    add_callout(doc, "最易混淆位", "考試附件叫 Appendix 1，但 course files 亦有 Appendix 1A／1B／1C。routing-core.md 要求兩者用唔同 namespace；exam Appendix 1 係 facts，course appendix 係教材。", fill=PALE_GOLD, border="8A6A18")

    add_heading(doc, "3. 逐步 trace：每一步睇咩 file、點解、再去邊", 1)

    steps = [
        ("Step 1｜由 workspace 規則入場", "AGENTS.md", "佢係呢個 folder 嘅入口指令，要求每一 turn 都讀 CLAUDE.md，亦要求每條問題先 consult Content.md。", "CLAUDE.md；之後 Content.md。", "確立本 folder 係 self-contained test bot，唔可以用 parent workspace 或外部法律填空。"),
        ("Step 2｜讀完整 operating brief", "CLAUDE.md", "佢定義考試結構、三種 answer modes、Routing v2 deterministic sequence、source hierarchy 同輸出格式。", "Content.md、routing-core.md、section-b.md、routing-v2/schema/route-plan.schema.json、routing-v2/scripts/validate_route_plan.py。", "知道 Q1 要先 routing、再 validate，最後先寫答案；KAP／gold／past answer 生成時禁止使用。"),
        ("Step 3｜讀考試題目本身", "PsLX_LZ-...Specimen-Paper-2...pdf（Q1：PDF page 2；exam Appendix 1：PDF pages 7-8）", "Q1 page 2 提供角色、jurisdiction 由 candidate 選、書面回覆要求同 20 marks；pages 7-8 提供 11 組 questionnaire 欄位、答案、疑問同 blanks。", "先將 facts 寫入 RoutePlan；再交畀 Content.md 做 issue routing。", "建立 18 個 fact records，例如 names、capital blank、client details gap、director gap、New York registered office、bank mandate blanks。"),
        ("Step 4｜判斷 answer mode", "CLAUDE.md + routing-core.md §3", "Q1 要『review and provide written response』，冇要求草擬 operative resolution／agreement，所以每個 answer unit 係 PROSE。", "section-b.md（Section B prose adapter）。", "RoutePlan.answer_unit.mode = PROSE；requested_document_chain.required = false。"),
        ("Step 5｜用 source map 找候選 files", "Content.md", "Content.md 係唯一 legal source map。佢唔直接講法律答案，而係按 issue 指向 Module／Appendix。Q1 同時有 formation、company type、capital、directors、secretary／agent、objects、nominee、audit、banking，所以 routes 要相加，唔可以只揀一個 Module。", "Modules 2、3、5、6、7、9、10、11；Appendix 6B；course Appendix 5 只作 check。", "產生 additive candidate-source union。"),
        ("Step 6｜做 facts ledger", "routing-core.md §4 + tmp/routeplans/question-1.json", "每個 material fact 必須有 disposition，避免一條 questionnaire query 被漏答。", "六個 locks；candidate routes；gaps handling。", "例如 New York registered office = used - outcome；client details only in New York = input gap；bank mandate blanks = input gap。"),
        ("Step 7｜鎖定六個核心問題", "routing-core.md §5 + question-1 RoutePlan", "如果 jurisdiction、entity、actor 或 lifecycle 未鎖定，就好容易混入另一司法管轄區或錯誤程序。", "BVI-specific passages in selected Modules and Appendix 6B。", "六 locks：BVI；BVI BC limited by shares；CSP administrator；pre-incorporation onboarding；before filing；BVI BCA + future memorandum/articles。"),
        ("Step 8｜第二次獨立 routing", "routing-core.md §7", "唔信第一次 topic match 就完事；要重新由 relationships 同 lifecycle 角度檢查有冇漏 route。", "Module 11（nominee／beneficial-owner control）、Module 9（registered agent）、Module 10（post-incorporation decisions）等 overlays。", "確認流程由 instructions／KYC 去到 incorporation、first director、shares、records、adviser、banking；冇要求 enforcement 或 winding-up。"),
        ("Step 9｜隔離 conditional／wrong routes", "routing-core.md §§8-10 + RoutePlan source_access", "要排除錯 jurisdiction、錯 constitutional model、KAP、prior answers，同埋對結果冇 unique contribution 嘅 sources。", "保留 mandatory／check-only sources；建立 forbidden_paths 同 prior_answer_paths。", "Appendix 1B／1C 被禁止；course Appendix 5 只 check field architecture；exam Appendix 1 先係 operative facts；XOR sets = 0。"),
        ("Step 10｜freeze allowlist，再開 source", "routing-core.md §§0、11 + RoutePlan source_access", "未開 substantive source 前先鎖死清單，防止寫到一半先搵有利材料補答案。每個 actual_open 都記 path、namespace、role、SHA-256。", "只開 allowlist 內 12 個 entries。", "allowlist_frozen = true；allowlist SHA-256 = e14d714e...ac7fb7。"),
        ("Step 11｜由每個 source 抽 answer building blocks", "已鎖定嘅 Modules／Appendices", "而家先讀 exact passages，逐個 questionnaire item 搵 unique contribution。", "final_trace 將每條 route 指去答案 1-11 嘅位置。", "例如 Module 7 → item 5 directors；Module 9 → items 6 and 9；Module 11 → items 4-6 and 9-11。"),
        ("Step 12｜按 schema 組成 RoutePlan", "routing-v2/schema/route-plan.schema.json", "schema 強制 RoutePlan 要有 answer_unit、namespaces、locks、facts、routes、source_access、gaps、final_trace、render_gate 等欄位，唔容許 answer 寫完先倒推。", "routing-v2/scripts/validate_route_plan.py。", "plan_id = Specimen2.Question1.BVI；render_gate 必須仍然係 not_rendered，證明答案未寫。"),
        ("Step 13｜validator 守門", "routing-v2/scripts/validate_route_plan.py + question-1-validation.json", "validator 唔讀自由文字答案，只驗 JSON schema 同跨欄位 invariants；失敗就唔可以 render。", "section-b.md（只喺 VALID + exit 0 之後）。", "15 invariants 全 PASS；status = VALID；plan SHA-256 = 9facc620...871f850。"),
        ("Step 14｜按 Section B 格式寫答案", "section-b.md", "佢規定答案要分 SUBMIT THIS 同 DO NOT SUBMIT check panel；亦要求 coverage、source、cross-check、validation、risk、confidence、verify。", "tmp/answers/question-1.md；再做 Done when checklist。", "用 questionnaire 1-11 做自然 headings，將 final_trace 嘅 building blocks 寫成完整 BVI advice。"),
    ]

    for title, file_name, why, nxt, result in steps:
        heading = add_heading(doc, title, 2)
        if title.startswith("Step 12"):
            heading.paragraph_format.page_break_before = True
        add_file_line(doc, file_name, why, nxt, result)

    doc.add_page_break()
    add_heading(doc, "4. Content.md 點樣將 Question 1 分流去各個 source", 1)
    add_text(doc, "下面係實際 issue-to-file map。重點係 routes 係 additive：同一條 questionnaire item 可以同時觸發兩至三個 Module，因為一個 Module 提供 substantive rule，另一個提供 procedure 或 service-provider overlay。")

    issue_rows = [
        ("1", "Name", "Module 3 §§3.1-3.6；Module 5 §§2.1、4", "name controls、availability／consent、constitutional name"),
        ("2", "Type of company", "Module 2 §§4.2、5；Module 3 §3.2；Appendix 6B", "limited by shares vs guarantee；BVI company form"),
        ("3", "Capital／shares", "Module 5 §§2.5、4.3；Module 6 §§1.2、4；Module 10 §5.9", "maximum shares、no-par／par choice、initial issue procedure"),
        ("4", "Shareholder／BO", "Module 3 onboarding；Module 11 §§2.3-2.10、3.1-3.6", "KYC、source of funds、nominee legal title vs beneficial ownership"),
        ("5", "Directors", "Module 7 §§3.1-3.3；Module 3 post-incorporation", "minimum number、corporate director、consent、first appointment"),
        ("6", "Secretary／RA", "Module 9 §§1.1-1.4、2.1-2.4；Module 11", "optional secretary、mandatory BVI registered agent、functions"),
        ("7", "Objects", "Module 5 §2.4", "BVI capacity、restricted-purpose alternative、need for business particulars"),
        ("8", "Articles", "Module 5 §§1、2、4", "standard articles as starting point；tailor director/quorum/share/control terms"),
        ("9", "Registered office", "Module 5 §2.2；Module 9 §2；Module 11", "must be in BVI；connect to licensed agent；New York only correspondence"),
        ("10", "Financial matters", "Module 3 §4；Module 10 §§5.9-5.10；Module 11", "audit/accounting、adviser appointment、year-end、records"),
        ("11", "Bank account", "Module 3 post-incorporation；Module 11 control/agency", "bank、currency、mandate、human signatories、board control"),
    ]
    add_table(doc, ["No.", "Questionnaire issue", "Content.md 指去", "對答案嘅 unique contribution"], issue_rows, [0.42, 1.25, 2.45, 2.38], font_size=8.05)

    add_heading(doc, "5. 六個 locks：點解最後揀 BVI", 1)
    add_text(doc, "Question 1 將 jurisdiction choice delegated 畀 candidate。routing-core.md 唔容許隨意揀；要比較 course coverage、actual-article fit、document-chain completeness 同 gaps。真正 tie 先按字母排序。呢次 BVI 路線喺 formation、company type、capital、directors、registered agent、objects、audit 同 banking 都有完整本地教材，所以一次過鎖定 BVI。")
    lock_rows = [
        ("Jurisdiction", "choice delegated", "British Virgin Islands", "本地 files 有完整 formation-to-administration coverage"),
        ("Regime/entity", "choice delegated", "BVI BC limited by shares", "Mr X 要 investment-holding，同 distributions／ownership 更相符"),
        ("Actor/capacity", "supplied", "CSP company administrator", "回覆 New York private banker，客戶係 Mr X"),
        ("Relationship", "supplied", "pre-incorporation onboarding", "未到成立後 transaction 或 litigation"),
        ("Lifecycle", "supplied", "before name approval/KYC/filing", "所以重點係 information gaps 同 formation decisions"),
        ("Governing instruments", "choice delegated", "BVI BCA + future M&A", "articles 未提供，只可以講要 tailor，唔可亂用 article number"),
    ]
    add_table(doc, ["Lock", "State", "Locked value", "點解"], lock_rows, [1.1, 1.05, 2.1, 2.25], font_size=8.3)

    add_heading(doc, "6. Relevance gate：點解有啲 file 唔睇／只係 check", 1)
    gate_rows = [
        ("Exam Appendix 1", "incorporated", "佢提供真正 questionnaire facts；係 requested deliverable 核心。"),
        ("Course Appendix 5", "checked-not-relevant", "只核對 questionnaire field architecture；唔可以取代 exam Appendix 1。"),
        ("Appendix 1B／1C", "forbidden", "實際 articles 未提供；用 Table A／English model 會污染 BVI-specific answer。"),
        ("Other jurisdictions", "excluded after lock", "BVI 一經選定，就唔可以混入 Bahamas／Jersey／Guernsey rules。"),
        ("KAP／gold answer", "forbidden during generation", "答案必須 blind；KAP 只可以喺答案 stage 同 hash-lock 之後作 marking。"),
        ("Prior Q1 artifacts", "prior-answer paths", "唔可用舊答案填 gap；RoutePlan 有明確記錄並排除。"),
        ("Drafting appendices", "not triggered", "Q1 只叫 written response，冇要求草擬 operative document；document chain = 0。"),
    ]
    add_table(doc, ["Source／route", "Verdict", "原因"], gate_rows, [2.0, 1.35, 3.15], font_size=8.4)

    add_heading(doc, "7. 三個 micro examples：一條 fact 點樣變成一句答案", 1)
    add_heading(doc, "Example A｜Registered office：New York 點解唔得", 2)
    add_text(doc, "Fact：exam Appendix 1 item 9 寫『Our office here in New York』。Disposition：used - outcome。Jurisdiction lock：BVI。Content route：Module 5（registered office substance）+ Module 9（registered agent）+ Module 11（service-provider context）。Exact-source check 後，答案先得出：registered office 要喺 BVI，通常由 licensed registered agent 提供；New York 只可做 correspondence address。")

    add_heading(doc, "Example B｜Capital：『maximum capital for minimum cost』點計", 2)
    add_text(doc, "Fact：authorised／issued capital 留白，banker 要 minimum cost。Content route：Module 5（memorandum/share terms）+ Module 6（authorised、issued、par/no-par）+ Module 10（initial issue procedure）。BVI lock 排除其他 jurisdiction fee/capital models。答案因此分開兩件事：公司可 authorise 50,000 no-par shares 以留喺最低 fee band；初步只 issue 一股畀 nominee for Mr X，其他留 unissued。")

    add_heading(doc, "Example C｜Client details：點解 New York file 唔夠", 2)
    add_text(doc, "Fact：item 4 只寫『Client details on file in New York』。Disposition：input gap。Content route：Module 3（formation/KYC）+ Module 11（beneficial owner/nominee control）。答案唔可以假設 CSP 已完成 due diligence，所以要求 Mr X 身份、地址、occupation、nationality、date/place of birth、source of funds、purpose、FATCA/CRS 等；nominee service 亦唔會令 Mr X 對 CSP 或 competent authorities 隱形。")

    add_heading(doc, "8. Frozen allowlist：實際獲准開嘅 files", 1)
    allow_rows = [
        ("1", "Exam PDF Q1", "exam_question", "角色、jurisdiction choice、20 marks、deliverable"),
        ("2", "Exam PDF Appendix 1", "exam_attachment", "items 1-11 facts、questions、blanks"),
        ("3", "Course-Manual-Module-02-Characteristics-of-a-Company.md", "course_manual", "company type、limited liability、asset-holding use"),
        ("4", "Course-Manual-Module-03-Company-Formation-and-Related-Issues.md", "course_manual", "formation、names、KYC、post-incorporation actions"),
        ("5", "Course-Manual-Module-05-The-Companys-Constitution.md", "course_manual", "M&A、objects/capacity、office、share terms"),
        ("6", "Course-Manual-Module-06-Equity-Capital-and-Distributions.md", "course_manual", "capital vs share authorisation/issue"),
        ("7", "Course-Manual-Module-07-Directors-Part-I-Role-Appointment-and-Removal-of-Directors.md", "course_manual", "director number、corporate director、appointment"),
        ("8", "Course-Manual-Module-09-Other-Officers-Secretary-and-Registered-Agent.md", "course_manual", "secretary and registered-agent rules/functions"),
        ("9", "Course-Manual-Module-10-Company-Decision-making-Procedures.md", "course_manual", "share issue、audit/accounting procedure"),
        ("10", "Course-Manual-Module-11-Company-Management-Services-and-Beneficial-Owner-Control.md", "course_manual", "nominee、BO control、adviser/bank arrangements"),
        ("11", "Appendix-6B-Types-of-company-available-in-selected-jurisdictions.md", "course_appendix", "BVI company forms cross-check"),
        ("12", "Appendix-5-New-Company-Questionnaire.md", "course_appendix / check-only", "field architecture only"),
    ]
    add_table(doc, ["#", "Path", "Namespace／role", "Unique contribution"], allow_rows, [0.35, 3.05, 1.35, 1.75], font_size=7.75)
    add_callout(doc, "Source lock", "allowlist freeze 後先 actual_open。每個 opened file 都記完整 SHA-256；答完先再由 final_trace 證明每個 mandatory route 真係喺答案有落點。", fill=PALE_BLUE, border=BLUE)

    add_heading(doc, "9. Validator 實際檢查咩", 1)
    invariant_rows = [
        ("Facts/claims", "每個 material fact 有 disposition，冇 orphan fact。"),
        ("Locks", "每個 lock 有 state/value/deciding fact；數量聲稱有 fact 支持。"),
        ("Namespaces", "exam question、exam attachment、course manual、course appendix 無混淆。"),
        ("Allowlist", "actual_open 必須喺 frozen list；forbidden/prior paths 冇被打開。"),
        ("Branches", "XOR route 只可選一條；每個 branch 有 deciding fact。Q1 無 XOR set。"),
        ("Document chain", "如果係 drafting，actor／signatory／components／attachments／execution 要對數。Q1 唔觸發。"),
        ("Gaps", "unresolved points 要 conditional／placeholder-backed；Q1 materials_gaps = 0。"),
        ("Final trace", "每個 selected route 只得一個 trace，並指明 answer location。"),
        ("Render gate", "RoutePlan 仲係 not_rendered，證明先驗證、後寫答案。"),
    ]
    add_table(doc, ["檢查組別", "Question 1 檢查內容"], invariant_rows, [1.45, 5.05], font_size=8.5)
    add_callout(doc, "驗證結果", "question-1-validation.json：exit code 0；15/15 invariants PASS；status VALID；plan SHA-256 9facc62047d03315adb7c00eab8b24c412d3c46924ef37d503f231735871f850。", fill=PALE_GREEN, border="2B6A45")

    add_heading(doc, "10. 最後點樣由 final trace 砌成答案", 1)
    trace_rows = [
        ("Preliminary paragraph", "Exam question + Appendix 6B", "State BVI once；apply consistently"),
        ("Answer 1", "Modules 3 + 5", "names、controlled words、banking risk、approval"),
        ("Answer 2", "Modules 2 + 3 + Appendix 6B", "BVI BC limited by shares vs guarantee"),
        ("Answer 3", "Modules 5 + 6 + 10", "50,000 shares、no-par、one initial share、unissued balance"),
        ("Answers 4-6", "Modules 3 + 7 + 9 + 11", "KYC、nominee、directors、secretary、registered agent"),
        ("Answers 7-9", "Modules 5 + 9 + 11", "objects/capacity、articles、BVI registered office"),
        ("Answers 10-11", "Modules 3 + 10 + 11", "audit/accounting、adviser、tax advice、year end、bank mandate"),
        ("Check panel", "section-b.md + validation report", "Coverage、authorities、sources、cross-check、risk、confidence、verify"),
    ]
    add_table(doc, ["答案位置", "由邊啲 route 餵入", "寫入嘅功能"], trace_rows, [1.35, 2.55, 2.6], font_size=8.4)

    add_heading(doc, "11. 呢套 routing 真正防止咗咩錯誤", 1)
    protection_rows = [
        ("混 jurisdiction", "BVI lock 之後排除 Bahamas/Jersey/Guernsey rules。"),
        ("漏答 questionnaire", "18 facts + dispositions，確保 items 1-11 同 blanks 都有處理。"),
        ("Exam appendix 撞 course appendix", "separate namespaces；exam Appendix 1 唔等於 course Appendix 1A-1C。"),
        ("先寫答案、後補 sources", "render_gate 必須 not_rendered；RoutePlan valid 先可以 prose。"),
        ("偷偷用 KAP／舊答案", "forbidden/prior paths + frozen allowlist + actual-open hashes。"),
        ("過度開 files", "hard relevance gate 要每個 source 有 unique contribution。"),
        ("將 prose 當 drafting", "mode = PROSE；document chain = 0；唔會無端草擬 resolution。"),
    ]
    add_table(doc, ["風險", "Routing v2 點防止"], protection_rows, [1.6, 4.9], font_size=8.7)

    add_heading(doc, "12. 你可以點樣自己重複呢個流程", 1)
    self_rows = [
        ("A", "先抄完整題目同附件，列出每個 fact、blank、command word、marks。"),
        ("B", "用 Content.md 對每個 issue 搵 base route + conditional overlays。"),
        ("C", "鎖六件事；唔知就標 unknown，唔好靠估。"),
        ("D", "做 relationship pass 同 lifecycle pass，睇有冇漏 route。"),
        ("E", "過 relevance gate；freeze allowlist；再開 exact passages。"),
        ("F", "填 RoutePlan JSON；validator 必須 exit 0／VALID。"),
        ("G", "最後先按 section adapter 寫答案同 coverage check。"),
    ]
    add_table(doc, ["次序", "動作"], self_rows, [0.55, 5.95], font_size=8.9)
    add_callout(doc, "判斷成功與否", "唔係『答案睇落合理』就算成功，而係：每個 fact 有 disposition、每個 source 有 unique contribution、每個 route 有 final trace、validator VALID、答案同 route 對得返。", fill=PALE_GOLD, border="8A6A18")

    add_heading(doc, "Appendix A｜主要工作 files 一覽", 1)
    appendix_rows = [
        ("AGENTS.md", "Workspace entry rule", "指去 CLAUDE.md；要求 Content.md first。"),
        ("CLAUDE.md", "Operating brief", "指去 Content、routing core、adapter、schema、validator。"),
        ("Content.md", "Declarative source map", "按 issue 指去 exact Modules／Appendices；本身唔係 law。"),
        ("routing-core.md", "Deterministic algorithm", "facts、locks、passes、relevance、allowlist、trace、gaps。"),
        ("section-b.md", "Section B renderer", "PROSE/DRAFTING handling、output blocks、coverage/style checklist。"),
        ("route-plan.schema.json", "Required JSON shape", "禁止缺 field 或事後倒推 route。"),
        ("validate_route_plan.py", "Machine validator", "輸出 VALID／INVALID、canonical plan hash、15 invariants。"),
        ("question-1.json", "Actual Q1 RoutePlan", "記錄 BVI locks、18 facts、12 allowlist entries、final trace。"),
        ("question-1-validation.json", "Validation evidence", "exit 0、status VALID、plan SHA。"),
        ("question-1.md", "Rendered answer", "submit-ready prose + non-submitted check panel。"),
    ]
    add_table(doc, ["File", "角色", "佢叫下一步做咩"], appendix_rows, [2.25, 1.45, 2.8], font_size=8.25)

    add_heading(doc, "Appendix B｜Audit identifiers", 1)
    audit_rows = [
        ("RoutePlan ID", "Specimen2.Question1.BVI"),
        ("Allowlist frozen", "true"),
        ("Allowlist SHA-256", "e14d714ec59aeb3ed87a697c1cf0d8169e00760a266b602b68a934f9caac7fb7"),
        ("Plan SHA-256", "9facc62047d03315adb7c00eab8b24c412d3c46924ef37d503f231735871f850"),
        ("Validation", "VALID；exit code 0；15 invariants PASS"),
        ("Materials gaps", "0（但 input gaps 有保留並喺答案要求補資料）"),
        ("XOR sets", "0"),
        ("Requested document chain", "required = false；counts all 0"),
    ]
    add_table(doc, ["Identifier", "Value"], audit_rows, [2.0, 4.5], font_size=8.4)

    add_heading(doc, "結語", 1)
    add_text(doc, "Question 1 嘅答案唔係由一個 Module 直接『搵出嚟』。佢係由 exam facts 開始，經 Content.md 將 issues 分流，再由 routing-core.md 合併、鎖定、過濾同驗證；最後 section-b.md 先將已驗證嘅 route 寫成答案。最重要嘅透明度係：任何一句主要 advice，都可以沿 final trace 行返去某條 fact、某個 route、某個 exact file。")

    props = doc.core_properties
    props.title = "Routing v2 Question 1 Step-by-Step Guide"
    props.subject = "Transparent worked example of deterministic company-law answer routing"
    props.author = "Codex"
    props.keywords = "Routing v2, STEP Company Law, Question 1, RoutePlan"

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_document())
